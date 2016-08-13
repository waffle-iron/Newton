import docx
import datetime

#from brain.models import StudentRoster, CurrentClass, Teacher


AMC_TESTS = {1: "Addition1Test.docx", 2: "Addition1Test.docx", 3: "Addition1Test.docx", 4: "Subtraction2Test.docx",}
AMC_FLASHCARDS = {}
# Temporary student name variable. Should be passed into the script eventually.
STUDENT_NAMES = {1: "Lucy Laika", 2: "Sam Smith",}


def get_closest_friday():
    today = datetime.date.today()
    if today.weekday() == 4:
        d = today
    elif today.weekday() != 4:
        if today.weekday() in [2, 3]:
            while today.weekday() != 4:
                today += datetime.timedelta(1)
            d = today
            # Add days until date is friday
        if today.weekday() in [5, 6, 0, 1]:
            while today.weekday() != 4:
                # Subtract days until date is Friday.
                today -= datetime.timedelta(1)
            d = today
    else:
        raise ValueError
    date_string = d.strftime("%A %B %d, %Y")
    return date_string


def write_documents():
    document = docx.Document(AMC_TESTS[4])
    today = get_closest_friday()
    for paragraph in document.paragraphs:
        if '<<student>>' in paragraph.text:
            paragraph.text = "Name: " + STUDENT_NAMES[1]
        if '<<date>>' in paragraph.text:
            paragraph.text = "Date: " + today

    document.save(STUDENT_NAMES[1]+'amctest.docx')

def search_and_replace():
    document = docx.Document(AMC_TESTS[4])
    today = get_closest_friday()
    document.add_heading('Heading, level 1', level=1)
    # for paragraph in document.paragraphs:
    #     if '<<student>>' in paragraph.text:
    #         current_paragraph = str(paragraph.text)
    #         current_paragraph.replace('<<student>>', STUDENT_NAMES[1])
    #         paragraph.text = current_paragraph
    #     if '<<date>>' in paragraph.text:
    #         current_paragraph = paragraph.text
    #         current_paragraph.replace('<<date>>', today)
    #         paragraph.text = current_paragraph

    document.save(STUDENT_NAMES[1]+'-amctest.docx')

def test_doc_create():
    doc = docx.Document('demo.docx')
    print(doc.paragraphs[0].text)
    # 'Document Title'
    print(doc.paragraphs[0].style)
    # 'Title'
    doc.paragraphs[0].style = 'Normal'
    print(doc.paragraphs[1].text)
    # 'A plain paragraph with some bold and some italic'
    # (doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
    #  doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
    # ('A plain paragraph with some ', 'bold', ' and some ', 'italic')
    doc.paragraphs[1].runs[0].style = 'QuoteChar'
    doc.paragraphs[1].runs[1].underline = True
    #doc.paragraphs[1].runs[3].underline = True
    doc.save('restyled.docx')

test_doc_create()