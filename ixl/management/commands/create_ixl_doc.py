from django.http import HttpResponse
import docx

from docx import Document
from docx.shared import Inches

def your_view(request):
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber')
    document.add_page_break()

    document.save('demo.docx')


    file_docx = open("./demo.docx", 'r')
    response = HttpResponse(mimetype='text/html')
    response['Content-Disposition'] = 'attachment; filename=IXLRecommendation.docx'
    response['Content-Encoding'] = 'UTF-8'
    response['Content-type'] = 'text/html; charset=UTF-8'
    response.write(file_docx.read())
    file_docx.close()
    return response